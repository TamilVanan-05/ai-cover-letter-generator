import os
import re
import json
import docx
import google.generativeai as genai
from config import Config

# Try to import pdfplumber for PDF extraction
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

def extract_text_from_pdf(file_path):
    """Extracts raw text from a PDF file using pdfplumber."""
    if not pdfplumber:
        raise ImportError("pdfplumber is not installed on this system.")
    
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()

def extract_text_from_docx(file_path):
    """Extracts raw text from a DOCX file using python-docx."""
    doc = docx.Document(file_path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
    return "\n".join(text).strip()

def parse_resume_text_with_gemini(text):
    """Uses Gemini API to semantically parse resume text into a structured profile JSON."""
    prompt = f"""
    You are an expert ATS Resume Parsing AI. Analyze the following raw resume text and extract the candidate's professional profile.
    
    RAW RESUME TEXT:
    {text}
    
    You MUST respond with ONLY a valid, raw JSON object. Do not wrap the JSON in markdown code blocks like ```json ... ``` or write any conversational text. Use exactly this JSON structure:
    {{
        "full_name": "Extract full name",
        "email": "Extract email address",
        "phone": "Extract phone number",
        "location": "Extract city, state or country",
        "linkedin": "Extract LinkedIn URL if present",
        "portfolio": "Extract personal website/portfolio URL if present",
        "experience_years": 5, // Integer estimate of years of experience
        "current_position": "Extract current or most recent job title",
        "previous_position": "Extract previous job title if present",
        "industry": "Determine the candidate's target industry sector",
        "skills": "Comma-separated list of technical/core skills",
        "certifications": "Comma-separated list of certifications if present",
        "education": "Markdown summary of educational degrees and colleges",
        "achievements": "List of key achievements, projects, or metrics (one per line, no bullet symbols, maximum 5)",
        "languages": "Comma-separated list of languages spoken if present",
        "projects": "Markdown summary of key personal/professional projects if present"
    }}
    """
    
    if Config.GEMINI_API_KEY:
        try:
            model = genai.GenerativeModel('gemini-1.5-flash')
            response = model.generate_content(prompt)
            if response and response.text:
                clean_text = response.text.strip()
                # Remove markdown wraps if the model ignores the prompt instruction
                if clean_text.startswith("```"):
                    lines = clean_text.splitlines()
                    if lines[0].startswith("```"):
                        lines = lines[1:]
                    if lines[-1].startswith("```"):
                        lines = lines[:-1]
                    clean_text = "\n".join(lines).strip()
                
                parsed_data = json.loads(clean_text)
                return parsed_data
        except Exception as e:
            print(f"Gemini resume parsing failed: {e}. Falling back to heuristic rules.")
            
    return parse_resume_text_heuristically(text)

def parse_resume_text_heuristically(text):
    """Heuristic rule-based regex parser fallback when LLM is unavailable."""
    profile = {
        "full_name": "",
        "email": "",
        "phone": "",
        "location": "",
        "linkedin": "",
        "portfolio": "",
        "experience_years": 0,
        "current_position": "",
        "previous_position": "",
        "industry": "Technology",
        "skills": "",
        "certifications": "",
        "education": "",
        "achievements": "",
        "languages": "",
        "projects": ""
    }
    
    # 1. Extract Email
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    if email_match:
        profile["email"] = email_match.group(0)
        
    # 2. Extract Phone
    phone_match = re.search(r'\(?\+?\d{1,3}\)?[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}', text)
    if phone_match:
        profile["phone"] = phone_match.group(0)
        
    # 3. Extract Name (often first line of resume)
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if lines:
        # Avoid picking up header labels
        candidate_name = lines[0]
        if len(candidate_name.split()) <= 4 and not any(kw in candidate_name.lower() for kw in ["resume", "curriculum", "vitae", "profile", "page"]):
            profile["full_name"] = candidate_name
        else:
            profile["full_name"] = "Applicant Name"
            
    # 4. Extract LinkedIn and Links
    linkedin_match = re.search(r'linkedin\.com/in/[\w-]+', text, re.IGNORECASE)
    if linkedin_match:
        profile["linkedin"] = "https://" + linkedin_match.group(0)
        
    urls = re.findall(r'https?://[^\s,]+', text)
    for url in urls:
        if "linkedin" not in url:
            profile["portfolio"] = url
            break
            
    # 5. Extract Years of Experience
    exp_matches = re.findall(r'(\d+)\+?\s*years?', text, re.IGNORECASE)
    if exp_matches:
        profile["experience_years"] = max([int(m) for m in exp_matches])
        
    # 6. Extract Skills (Common Skill Sets Matching)
    common_skills = [
        "python", "javascript", "react", "vue", "angular", "node", "django", "flask",
        "sql", "nosql", "mongodb", "postgresql", "mysql", "aws", "azure", "docker", "kubernetes",
        "git", "ci/cd", "agile", "scrum", "project management", "product management", "ui/ux",
        "figma", "html", "css", "java", "c++", "c#", "go", "rust", "php", "ruby", "machine learning",
        "deep learning", "nlp", "data science", "excel", "seo", "salesforce", "marketing"
    ]
    found_skills = []
    for skill in common_skills:
        if re.search(r'\b' + re.escape(skill) + r'\b', text, re.IGNORECASE):
            found_skills.append(skill.title())
    profile["skills"] = ", ".join(found_skills)
    
    # 7. Extract Positions
    job_titles = ["software engineer", "developer", "product manager", "project manager", "designer", "analyst", "consultant", "administrator", "coordinator"]
    found_positions = []
    for title in job_titles:
        matches = re.findall(r'\b' + re.escape(title) + r'\b', text, re.IGNORECASE)
        if matches:
            found_positions.append(title.title())
    if found_positions:
        profile["current_position"] = found_positions[0]
        if len(found_positions) > 1:
            profile["previous_position"] = found_positions[1]
            
    # 8. Industry mapping
    if any(tech in profile["skills"].lower() for tech in ["python", "javascript", "react", "sql", "aws", "docker"]):
        profile["industry"] = "Technology / Software"
        
    # 9. Simple Education parsing
    edu_keywords = ["university", "college", "institute", "school", "bachelor", "master", "phd", "b.s.", "m.s.", "b.tech", "m.tech", "diploma"]
    edu_lines = []
    for line in lines:
        if any(re.search(r'\b' + re.escape(kw) + r'\b', line, re.IGNORECASE) for kw in edu_keywords):
            edu_lines.append(line)
    if edu_lines:
        profile["education"] = "\n".join(edu_lines[:3])
        
    # 10. Achievements
    ach_keywords = ["optimized", "delivered", "increased", "decreased", "saved", "reduced", "led", "managed", "spearheaded", "designed", "implemented", "achieved"]
    ach_lines = []
    for line in lines:
        # Check if line contains a metric (percentage or numbers) and an action verb
        if any(char.isdigit() for char in line) and any(re.search(r'\b' + re.escape(kw) + r'\b', line, re.IGNORECASE) for kw in ach_keywords):
            ach_lines.append(line)
            if len(ach_lines) >= 3:
                break
    profile["achievements"] = "\n".join(ach_lines) if ach_lines else "Successfully led key project executions and delivered software engineering components."
    
    return profile

def parse_resume_file(file_path, file_extension):
    """Main function to parse resume files (.pdf or .docx) and return structured dictionary."""
    ext = file_extension.lower().strip('.')
    raw_text = ""
    
    if ext == 'pdf':
        raw_text = extract_text_from_pdf(file_path)
    elif ext in ['docx', 'doc']:
        raw_text = extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file extension: .{ext}")
        
    if not raw_text:
        raise ValueError("Could not extract any readable text content from the uploaded resume file.")
        
    return parse_resume_text_with_gemini(raw_text)
