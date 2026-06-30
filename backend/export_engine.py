import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx_file(cover_letter_data):
    """Generates a professional DOCX cover letter using python-docx with styled typography."""
    
    doc = Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Configure document style defaults
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Charcoal
    
    # 1. Header (Candidate Contact Info)
    # Highlight candidate name in bold and primary color
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(cover_letter_data.get('full_name', 'Applicant Name').upper())
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Slate 900
    
    # Sub-header contact line
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(24)
    
    contact_items = []
    if cover_letter_data.get('location'):
        contact_items.append(cover_letter_data.get('location'))
    if cover_letter_data.get('phone'):
        contact_items.append(cover_letter_data.get('phone'))
    if cover_letter_data.get('email'):
        contact_items.append(cover_letter_data.get('email'))
    if cover_letter_data.get('linkedin'):
        contact_items.append("LinkedIn")
        
    contact_text = "  |  ".join(contact_items)
    contact_run = contact_p.add_run(contact_text)
    contact_run.font.size = Pt(9.5)
    contact_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B) # Slate 500
    
    # 2. Date
    date_p = doc.add_paragraph()
    date_p.paragraph_format.space_after = Pt(12)
    # Convert date format if needed, or print custom date
    import datetime
    current_date = datetime.date.today().strftime("%B %d, %Y")
    date_run = date_p.add_run(current_date)
    date_run.font.size = Pt(11)
    
    # 3. Recipient Address Block
    rec_p = doc.add_paragraph()
    rec_p.paragraph_format.space_after = Pt(18)
    
    hiring_mgr = cover_letter_data.get('hiring_manager', '')
    company = cover_letter_data.get('company_name', '')
    job_loc = cover_letter_data.get('job_location', '')
    
    mgr_line = hiring_mgr if hiring_mgr and hiring_mgr != "Hiring Manager" else "Hiring Committee"
    rec_lines = [mgr_line, company]
    if job_loc:
        rec_lines.append(job_loc)
        
    rec_run = rec_p.add_run("\n".join(rec_lines))
    rec_run.font.size = Pt(11)
    rec_run.font.bold = True
    
    # 4. Content split by paragraph
    content = cover_letter_data.get('content', '')
    paragraphs_list = [p.strip() for p in content.split('\n\n') if p.strip()]
    
    # If the AI letter already includes the greeting/signoff, we should be careful not to double add,
    # but normally the cover letter object 'content' has everything (greeting, body, sign-off).
    # We will loop and print each paragraph cleanly.
    
    first_p = True
    for p_text in paragraphs_list:
        # Skip if paragraph is just double name or matches salutation (we want to check formatting)
        # Check if this paragraph is the salutation
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(12)
        
        run = p.add_run(p_text)
        run.font.size = Pt(11)
        
        # If it's a signature or contact line, format it slightly smaller
        if "email:" in p_text.lower() or "phone:" in p_text.lower():
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69) # Slate 600
            run.font.size = Pt(10)
            
    # Save document to memory stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def send_cover_letter_email(email_recipient, subject, letter_content):
    """Simulates sending the cover letter to a user's inbox or recruiter email."""
    # SMTPServer configuration template (Mocked for testing env)
    try:
        # In a real production SaaS, you would use:
        # import smtplib
        # from email.mime.text import MIMEText
        # msg = MIMEText(letter_content)
        # msg['Subject'] = subject
        # msg['From'] = "noreply@aicovergenerator.com"
        # msg['To'] = email_recipient
        # ... smtplib SMTP login and send
        
        print(f"[SMTP Simulator] Sending email to {email_recipient} with subject: {subject}")
        return True, "Email sent successfully (Simulated SMTP)"
    except Exception as e:
        return False, str(e)
